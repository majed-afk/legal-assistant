"""
Contract analysis service.
Extracts text from uploaded contracts (PDF/DOCX/plain text),
detects contract type, and streams a legal analysis using Claude.
"""
from __future__ import annotations

import io
import logging
import re
import time
from typing import Generator, Optional

import anthropic

from backend.config import ANTHROPIC_API_KEY, CLAUDE_MODEL

log = logging.getLogger("sanad.contract_analyzer")


# ══════════════════════════════════════════════════════════════
# System prompt for contract analysis
# ══════════════════════════════════════════════════════════════

SYSTEM_PROMPT_CONTRACT = """أنت سند — محلل عقود قانوني سعودي متخصص.
مهمتك تحليل العقود المقدمة مقابل الأنظمة السعودية المعمول بها.

## تخصصاتك (5 أنظمة + ضوابط)
1. **نظام المعاملات المدنية** — العقود والالتزامات: بيع، إيجار، مقاولة، وكالة، كفالة، شركة، مضاربة، ملكية، تعويض
2. **نظام الأحوال الشخصية** — الزواج والوصية والهبة
3. **نظام الإثبات** — طرق إثبات العقود: كتابة، شهادة، إقرار، دليل رقمي + ضوابط الإثبات إلكترونياً
4. **نظام المرافعات الشرعية** — الإجراءات والاختصاص
5. **نظام المحاكم التجارية** — اختصاص المحكمة التجارية، الدعاوى التجارية، أوامر الأداء، الطلبات المستعجلة، الاستئناف والنقض (مهم للعقود التجارية)

## مهامك عند تحليل العقد:
1. **تصنيف العقد** — حدد نوعه (بيع، إيجار، مقاولة، وكالة، كفالة، شركة، مضاربة، صلح، هبة...)
2. **فحص الأركان الأساسية** — هل توفرت أركان العقد (رضا، محل، سبب) وفق نظام المعاملات المدنية؟
3. **البنود المخالفة** — هل يوجد بنود تخالف النظام؟ (شروط باطلة، غرر، غبن فاحش، شرط جزائي مبالغ فيه...)
4. **البنود الناقصة** — بنود مهمة يجب أن تكون موجودة ولكنها غائبة
5. **تقييم المخاطر** — بنود قد تضر أحد الأطراف أو تحتمل تفسيرات متعددة
6. **التوصيات** — اقتراحات لتحسين العقد وحماية الأطراف

## قواعد إلزامية:
- استند حصرياً للمواد النظامية المرفقة — لا تستخدم معرفتك العامة
- كل ملاحظة يجب أن تكون مسنودة بـ: رقم المادة + اسم النظام
- لا تذكر أي رقم مادة غير موجود في النصوص المرفقة
- صنّف كل ملاحظة بأيقونة: 🔴 مخالفة | 🟡 تحذير | 🟢 توصية
- استخدم الفصحى الواضحة البسيطة
- إذا لم تجد ما يكفي من المواد قل ذلك بوضوح

## هيكل التقرير:
### 📋 تصنيف العقد
(نوع العقد، الأطراف، الموضوع)

### ✅ الأركان الأساسية
(هل توفر الرضا والمحل والسبب؟)

### 🔍 تحليل البنود
(تحليل كل بند مهم مع المواد المرتبطة)

### ⚠️ المخاطر والتحذيرات
(بنود مخالفة أو خطيرة مع التصنيف 🔴🟡)

### 📝 التوصيات
(اقتراحات تحسين مع التصنيف 🟢)

### 📚 المواد النظامية المرتبطة
(قائمة المواد المستخدمة في التحليل)

⚖️ هذا تحليل أولي لا يُغني عن مراجعة محامي مرخص."""


# ══════════════════════════════════════════════════════════════
# Text extraction
# ══════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        text = "\n".join(text_parts)
        if not text.strip():
            raise ValueError("لم يتم استخراج نص من ملف PDF — قد يكون الملف صورة (scanned)")
        return text
    except ImportError:
        raise ValueError("مكتبة PyPDF2 غير متوفرة")
    except Exception as e:
        if "استخراج" in str(e):
            raise
        raise ValueError(f"خطأ في قراءة ملف PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        text = "\n".join(text_parts)
        if not text.strip():
            raise ValueError("لم يتم استخراج نص من ملف Word — الملف فارغ")
        return text
    except ImportError:
        raise ValueError("مكتبة python-docx غير متوفرة")
    except Exception as e:
        if "استخراج" in str(e) or "فارغ" in str(e):
            raise
        raise ValueError(f"خطأ في قراءة ملف Word: {str(e)}")


# ══════════════════════════════════════════════════════════════
# Contract type detection
# ══════════════════════════════════════════════════════════════

# Keyword patterns for contract type detection
CONTRACT_TYPE_PATTERNS = {
    "بيع": [
        r"عقد\s*بيع", r"البائع", r"المشتري", r"ثمن\s*البيع",
        r"المبيع", r"نقل\s*الملكية", r"تسليم\s*المبيع",
    ],
    "إيجار": [
        r"عقد\s*إيجار", r"عقد\s*ايجار", r"المؤجر", r"المستأجر",
        r"الأجرة", r"مدة\s*الإيجار", r"العين\s*المؤجرة",
    ],
    "مقاولة": [
        r"عقد\s*مقاولة", r"المقاول", r"صاحب\s*العمل",
        r"الأعمال\s*المطلوبة", r"مدة\s*التنفيذ", r"مستخلصات",
    ],
    "وكالة": [
        r"عقد\s*وكالة", r"الموكل", r"الوكيل", r"التوكيل",
        r"نطاق\s*الوكالة", r"صلاحيات",
    ],
    "شركة": [
        r"عقد\s*شركة", r"عقد\s*تأسيس", r"الشريك", r"الشركاء",
        r"رأس\s*المال", r"حصة", r"الأرباح\s*والخسائر",
    ],
    "كفالة": [
        r"عقد\s*كفالة", r"الكفيل", r"المكفول", r"الدائن",
        r"المدين", r"ضمان",
    ],
    "مضاربة": [
        r"عقد\s*مضاربة", r"رب\s*المال", r"المضارب",
        r"رأس\s*مال\s*المضاربة",
    ],
    "هبة": [
        r"عقد\s*هبة", r"الواهب", r"الموهوب\s*له",
    ],
    "صلح": [
        r"عقد\s*صلح", r"تسوية", r"اتفاقية\s*صلح",
        r"الطرف\s*الأول.*الطرف\s*الثاني.*تنازل",
    ],
    "عمل": [
        r"عقد\s*عمل", r"عقد\s*توظيف", r"الموظف", r"صاحب\s*العمل",
        r"الراتب", r"مدة\s*العقد.*سنة", r"فترة\s*التجربة",
    ],
}


def detect_contract_type(text: str) -> str:
    """Detect contract type from text using keyword patterns."""
    scores: dict[str, int] = {}
    text_lower = text[:3000]  # Only check first 3000 chars for efficiency

    for contract_type, patterns in CONTRACT_TYPE_PATTERNS.items():
        score = 0
        for pattern in patterns:
            matches = re.findall(pattern, text_lower)
            score += len(matches)
        if score > 0:
            scores[contract_type] = score

    if not scores:
        return "عام"

    return max(scores, key=scores.get)


# Topic mapping: contract type → RAG query topics
CONTRACT_RAG_QUERIES = {
    "بيع": "عقد البيع: أركان البيع، التزامات البائع، التزامات المشتري، ضمان العيب، تسليم المبيع، ثمن البيع",
    "إيجار": "عقد الإيجار: التزامات المؤجر، التزامات المستأجر، الأجرة، مدة الإيجار، إنهاء العقد، صيانة العين المؤجرة",
    "مقاولة": "عقد المقاولة: التزامات المقاول، التزامات صاحب العمل، تسليم الأعمال، ضمان العيوب، فسخ العقد",
    "وكالة": "عقد الوكالة: حدود الوكالة، التزامات الوكيل، التزامات الموكل، انتهاء الوكالة، مسؤولية الوكيل",
    "شركة": "عقد الشركة: أركان الشركة، حصص الشركاء، إدارة الشركة، توزيع الأرباح والخسائر، انقضاء الشركة",
    "كفالة": "عقد الكفالة: أركان الكفالة، التزامات الكفيل، حدود الكفالة، انقضاء الكفالة",
    "مضاربة": "عقد المضاربة: شروط المضاربة، التزامات المضارب، توزيع الربح، انتهاء المضاربة",
    "هبة": "عقد الهبة: أركان الهبة، شروط الهبة، الرجوع في الهبة",
    "صلح": "عقد الصلح: أركان الصلح، آثار الصلح، بطلان الصلح",
    "عمل": "عقد العمل: التزامات العامل، التزامات صاحب العمل، إنهاء العقد، مكافأة نهاية الخدمة",
    "عام": "العقود: أركان العقد، شروط صحة العقد، آثار العقد، بطلان العقد، فسخ العقد",
}


# ══════════════════════════════════════════════════════════════
# Claude streaming analysis
# ══════════════════════════════════════════════════════════════

def _get_client() -> anthropic.Anthropic:
    if not ANTHROPIC_API_KEY:
        raise ValueError("ANTHROPIC_API_KEY غير مُعَد")
    return anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)


def stream_contract_analysis(
    contract_text: str,
    context: str,
    contract_type: str = "عام",
) -> Generator[str, None, None]:
    """Stream contract analysis token-by-token using Claude API."""
    client = _get_client()

    # Truncate very long contracts to avoid token limits
    max_contract_chars = 12000
    if len(contract_text) > max_contract_chars:
        contract_text = contract_text[:max_contract_chars] + "\n\n... [تم اقتطاع بقية العقد — يُرجى تقسيم العقود الطويلة]"

    user_message = f"""## نص العقد المطلوب تحليله:
{contract_text}

---

## نوع العقد المُكتشف: {contract_type}

---

## 📚 المواد النظامية المسترجعة:
{context}

---

⛔ حلّل العقد أعلاه استناداً حصرياً للمواد النظامية المرفقة. لا تذكر مواد غير مقدمة لك."""

    messages = [{"role": "user", "content": user_message}]

    max_retries = 3
    for attempt in range(max_retries):
        try:
            with client.messages.stream(
                model=CLAUDE_MODEL,
                max_tokens=4000,
                system=SYSTEM_PROMPT_CONTRACT,
                messages=messages,
            ) as stream:
                for text in stream.text_stream:
                    yield text
            return
        except anthropic.RateLimitError:
            if attempt < max_retries - 1:
                wait = 2 ** attempt * 5
                log.warning("Rate limited, retrying in %ds (attempt %d)", wait, attempt + 1)
                time.sleep(wait)
            else:
                raise
        except anthropic.APIStatusError as e:
            if e.status_code == 529 and attempt < max_retries - 1:
                wait = 2 ** attempt * 5
                log.warning("API overloaded (529), retrying in %ds", wait)
                time.sleep(wait)
            else:
                raise
